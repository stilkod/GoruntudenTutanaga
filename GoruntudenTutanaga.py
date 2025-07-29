import tkinter as tk
from tkinter import filedialog, messagebox
import vlc
from docx import Document
from docx.shared import Inches
import os, time
from PIL import Image, ImageTk, ImageDraw, ImageFont

def format_time(seconds):
    """Saniye cinsinden alınan zamanı hh:mm:ss ya da mm:ss formatında döndürür."""
    seconds = int(seconds)
    hrs = seconds // 3600
    mins = (seconds % 3600) // 60
    secs = seconds % 60
    if hrs > 0:
        return f"{hrs:02d}:{mins:02d}:{secs:02d}"
    else:
        return f"{mins:02d}:{secs:02d}"

class VideoAnnotationApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Görüntü İzleme ve Tutanak Uygulaması")
        self.geometry("1024x768")
        
        vlc_options = [
            '--file-caching=3000',
            '--network-caching=3000',
            '--no-video-title-show',
            '--avcodec-hw=none',
        ]
        self.vlc_instance = vlc.Instance(vlc_options)
        self.player = self.vlc_instance.media_player_new()
        
        self.video_panel = tk.Frame(self, bg="black")
        self.video_panel.pack(fill=tk.BOTH, expand=True)
        
        self.slider_dragging = False
        self.event_list = []
        self.current_video = None
        
        controls_frame = tk.Frame(self)
        controls_frame.pack(fill=tk.X, padx=5, pady=5)
        
        self.open_button = tk.Button(controls_frame, text="Video Aç", command=self.open_video)
        self.open_button.pack(side=tk.LEFT, padx=2)
        self.play_button = tk.Button(controls_frame, text="Oynat", command=self.play_video)
        self.play_button.pack(side=tk.LEFT, padx=2)
        self.pause_button = tk.Button(controls_frame, text="Duraklat", command=self.pause_video)
        self.pause_button.pack(side=tk.LEFT, padx=2)
        self.skip_backward_button = tk.Button(controls_frame, text="Geri Sar (←)", command=self.skip_backward)
        self.skip_backward_button.pack(side=tk.LEFT, padx=2)
        self.skip_forward_button = tk.Button(controls_frame, text="İleri Sar (→)", command=self.skip_forward)
        self.skip_forward_button.pack(side=tk.LEFT, padx=2)
        
        tk.Label(controls_frame, text="Tespit Edilen Durum:").pack(side=tk.LEFT, padx=5)
        self.event_entry = tk.Entry(controls_frame, width=40)
        self.event_entry.pack(side=tk.LEFT, padx=2, fill=tk.X, expand=True)
        self.event_entry.bind("<Return>", lambda event: self.add_text_event())
        
        tk.Button(controls_frame, text="Tespit Ekle", command=self.add_text_event).pack(side=tk.LEFT, padx=2)
        tk.Button(controls_frame, text="İşaretçi Ekle", command=self.add_pointer_event).pack(side=tk.LEFT, padx=2)
        tk.Button(controls_frame, text="Rapor Oluştur", command=self.create_report).pack(side=tk.LEFT, padx=2)
        tk.Button(controls_frame, text="Tespit Sil", command=self.delete_event).pack(side=tk.LEFT, padx=2)
        
        self.events_listbox = tk.Listbox(self, height=8)
        self.events_listbox.pack(fill=tk.X, padx=5, pady=5)
        
        self.time_slider = tk.Scale(self, from_=0, to=1000, orient=tk.HORIZONTAL, showvalue=0)
        self.time_slider.pack(fill=tk.X, padx=5, pady=2)

        self.time_slider.bind("<ButtonPress-1>", self.on_slider_press)
        self.time_slider.bind("<ButtonRelease-1>", self.on_slider_release)

        self.time_label = tk.Label(self, text="00:00 / 00:00")
        self.time_label.pack(pady=2)
        
        self.is_playing = False
        self.video_loaded = False
        
        self.bind("<space>", self.toggle_play_pause)
        self.bind("<Right>", self.on_right_arrow)
        self.bind("<Left>", self.on_left_arrow)
        self.update_time()
        self.messagebox_options = {"parent": self}

    def on_slider_press(self, event):
        self.slider_dragging = True

    def on_slider_release(self, event):
        self.slider_dragging = False
        if self.video_loaded:
            slider_value = self.time_slider.get()
            self.player.set_position(slider_value / 1000.0)

    def set_video_panel(self):
        if os.name == "nt":
            self.player.set_hwnd(self.video_panel.winfo_id())
        elif os.name == "posix":
            if os.uname().sysname == "Darwin":
                self.player.set_nsobject(self.video_panel.winfo_id())
            else:
                self.player.set_xwindow(self.video_panel.winfo_id())

    def open_video(self):
        video_path = filedialog.askopenfilename(
            title="Video Seç", 
            filetypes=[("Video Dosyaları", "*.mp4 *.avi *.mkv"), ("Tüm Dosyalar", "*.*")]
        )
        if video_path:
            self.current_video = video_path
            media = self.vlc_instance.media_new(video_path)
            self.player.set_media(media)
            self.video_loaded = True
            
            self.set_video_panel()
            self.play_video()
            self.time_slider.set(0)

    def play_video(self):
        if self.video_loaded and not self.player.is_playing():
            self.player.play()
            self.is_playing = True

    def pause_video(self):
        if self.video_loaded and self.player.is_playing():
            self.player.pause()
            self.is_playing = False

    def skip_forward(self, ms=5000):
        if self.video_loaded:
            self.player.set_time(self.player.get_time() + ms)

    def skip_backward(self, ms=5000):
        if self.video_loaded:
            self.player.set_time(self.player.get_time() - ms)

    def add_text_event(self):
        if not self.video_loaded:
            messagebox.showwarning("Uyarı", "Lütfen önce bir video dosyası açın.", **self.messagebox_options)
            return

        current_time_ms = self.player.get_time()
        current_time_sec = current_time_ms / 1000.0
        event_text = self.event_entry.get().strip()
        if not event_text:
            messagebox.showwarning("Uyarı", "Lütfen tespit edilen durumu girin.", **self.messagebox_options)
            return
        
        was_playing = self.is_playing
        if was_playing:
            self.pause_video()
            time.sleep(0.1)

        timestamp = time.time()
        screenshot_filename = f"ekrangoruntusu_{timestamp:.2f}.png"
        result = self.player.video_take_snapshot(0, screenshot_filename, 0, 0)

        if was_playing:
            self.play_video()

        if result != 0:
            messagebox.showerror("Hata", "Ekran görüntüsü alınamadı!", **self.messagebox_options)
            return
        
        event_info = {
            "type": "screenshot",
            "time": current_time_sec,
            "description": event_text,
            "image": screenshot_filename,
            "video": self.current_video
        }
        self.event_list.append(event_info)
        display_text = f"Tespit {len(self.event_list)} - Zaman: {format_time(current_time_sec)} - {event_text}"
        self.events_listbox.insert(tk.END, display_text)
        self.event_entry.delete(0, tk.END)

    def add_pointer_event(self):
        if not self.video_loaded:
            messagebox.showwarning("Uyarı", "Lütfen önce bir video dosyası açın.", **self.messagebox_options)
            return
        
        was_playing = self.is_playing
        self.pause_video()
        time.sleep(0.1)

        current_time_ms = self.player.get_time()
        current_time_sec = current_time_ms / 1000.0
        
        timestamp = time.time()
        screenshot_filename = f"ekrangoruntusu_temp_{timestamp:.2f}.png"
        result = self.player.video_take_snapshot(0, screenshot_filename, 0, 0)
        if result != 0 or not os.path.exists(screenshot_filename):
            messagebox.showerror("Hata", "Ekran görüntüsü alınamadı!", **self.messagebox_options)
            if was_playing:
                self.play_video()
            return
        
        AnnotationWindow(self, screenshot_filename, current_time_sec, was_playing)

    def create_report(self):
        if not self.event_list:
            messagebox.showinfo("Bilgi", "Eklenmiş tespit bilgisi bulunmamaktadır.", **self.messagebox_options)
            return
        
        template_path = "RPSABLON.docx"
        if not os.path.exists(template_path):
            messagebox.showerror("Hata", f"Şablon dosyası bulunamadı! '{template_path}'", **self.messagebox_options)
            return
        
        try:
            doc = Document(template_path)
        except Exception as e:
            messagebox.showerror("Hata", f"Şablon dosyası yüklenemedi: {e}", **self.messagebox_options)
            return
        
        videos = {}
        for event in self.event_list:
            vid = event.get("video", "Bilinmeyen Video")
            video_name = os.path.basename(vid) if vid else "Bilinmeyen Video"
            videos.setdefault(video_name, []).append(event)
        
        images_to_delete = []

        for video_file, events in videos.items():
            doc.add_heading(f"Video: {video_file}", level=1)
            for idx, event in enumerate(events, start=1):
                doc.add_heading(f"Tespit {idx}", level=2)
                doc.add_paragraph(f"Zaman: {format_time(event['time'])}")
                if event.get("description"):
                    doc.add_paragraph(f"Açıklama: {event['description']}")
                
                image_path = event.get("image")
                if image_path and os.path.exists(image_path):
                    try:
                        doc.add_picture(image_path, width=Inches(6.0))
                        images_to_delete.append(image_path)
                    except Exception as e:
                        doc.add_paragraph(f"[Resim eklenemedi: {e}]")
                else:
                    doc.add_paragraph("[Resim dosyası bulunamadı.]")
                doc.add_paragraph()

        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Dosyası", "*.docx")],
            title="Raporu Kaydet"
        )
        
        if save_path:
            try:
                doc.save(save_path)
                messagebox.showinfo("Başarılı", "Rapor başarıyla kaydedildi.", **self.messagebox_options)

                # Rapor kaydedildikten sonra resimleri sil ve listeleri temizle
                for image_path in images_to_delete:
                    try:
                        os.remove(image_path)
                    except OSError as e:
                        print(f"Hata: Resim dosyası silinemedi {image_path}: {e}")
                
                self.event_list.clear()
                self.events_listbox.delete(0, tk.END)
                messagebox.showinfo("Temizlik", "Kullanılan ekran görüntüleri silindi ve liste temizlendi.", **self.messagebox_options)

            except Exception as e:
                messagebox.showerror("Hata", f"Rapor kaydedilirken veya temizlik yapılırken hata oluştu:\n{e}", **self.messagebox_options)

    def delete_event(self):
        selected_indices = self.events_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("Uyarı", "Lütfen silmek istediğiniz tespiti seçin.", **self.messagebox_options)
            return
        
        for index in reversed(selected_indices):
            event_to_delete = self.event_list[index]
            if os.path.exists(event_to_delete['image']):
                try:
                    os.remove(event_to_delete['image'])
                except OSError as e:
                    print(f"Error deleting image file {event_to_delete['image']}: {e}")

            del self.event_list[index]
            self.events_listbox.delete(index)

    def update_time(self):
        if self.video_loaded:
            length_ms = self.player.get_length()
            if length_ms > 0:
                current_time_ms = self.player.get_time()
                
                if not self.slider_dragging:
                    position = self.player.get_position()
                    self.time_slider.set(int(position * 1000))
                
                current_time_sec = current_time_ms / 1000
                length_sec = length_ms / 1000
                self.time_label.config(text=f"{format_time(current_time_sec)} / {format_time(length_sec)}")

        self.after(250, self.update_time)

    def toggle_play_pause(self, event=None):
        if self.focus_get() == self.event_entry:
            return
        if self.is_playing:
            self.pause_video()
        else:
            self.play_video()

    def on_right_arrow(self, event=None):
        if self.focus_get() != self.event_entry:
            self.skip_forward()

    def on_left_arrow(self, event=None):
        if self.focus_get() != self.event_entry:
            self.skip_backward()

class AnnotationWindow(tk.Toplevel):
    def __init__(self, parent, screenshot_path, current_time_sec, was_playing):
        super().__init__(parent)
        self.parent = parent
        self.screenshot_path = screenshot_path
        self.current_time_sec = current_time_sec
        self.was_playing = was_playing
        self.title("İşaretçi Anotasyonu")
        self.transient(parent)
        self.grab_set()

        self.protocol("WM_DELETE_WINDOW", self.on_close)
        
        try:
            self.image = Image.open(screenshot_path)
        except FileNotFoundError:
            messagebox.showerror("Hata", "Ekran görüntüsü dosyası bulunamadı.", parent=self)
            self.destroy()
            return

        video_panel_width = self.parent.video_panel.winfo_width()
        video_panel_height = self.parent.video_panel.winfo_height()
        img_w, img_h = self.image.size
        
        aspect_ratio = img_h / img_w
        new_width = video_panel_width
        new_height = int(new_width * aspect_ratio)

        if new_height > video_panel_height:
            new_height = video_panel_height
            new_width = int(new_height / aspect_ratio)
        
        self.image = self.image.resize((new_width, new_height), Image.Resampling.LANCZOS)
        
        self.annotated_image = self.image.copy()
        self.draw = ImageDraw.Draw(self.annotated_image)
        
        self.canvas = tk.Canvas(self, width=self.image.width, height=self.image.height, cursor="cross")
        self.canvas.pack()
        self.photo = ImageTk.PhotoImage(self.image)
        self.canvas_image = self.canvas.create_image(0, 0, anchor="nw", image=self.photo)
        
        self.start_x = None
        self.start_y = None
        self.arrow_end = None
        self.current_arrow_id = None
        self.text_entry = None
        self.annotations = []

        self.canvas.bind("<ButtonPress-1>", self.on_button_press)
        self.canvas.bind("<B1-Motion>", self.on_mouse_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_button_release)
        
        button_frame = tk.Frame(self)
        button_frame.pack(pady=5, fill=tk.X, padx=10)
        
        tk.Button(button_frame, text="Kaydet ve Ekle", command=self.save_annotation).pack(side=tk.RIGHT)
        tk.Button(button_frame, text="Sıfırla", command=self.reset_annotations).pack(side=tk.RIGHT, padx=5)
        self.messagebox_options = {"parent": self}

    def reset_annotations(self):
        self.canvas.delete("all")
        self.annotated_image = self.image.copy()
        self.draw = ImageDraw.Draw(self.annotated_image)
        self.photo = ImageTk.PhotoImage(self.image)
        self.canvas_image = self.canvas.create_image(0, 0, anchor="nw", image=self.photo)
        self.annotations.clear()

    def on_button_press(self, event):
        if self.text_entry: return
        self.start_x = event.x
        self.start_y = event.y
        self.current_arrow_id = self.canvas.create_line(self.start_x, self.start_y, event.x, event.y,
                                                  arrow=tk.LAST, fill="red", width=3)
    
    def on_mouse_drag(self, event):
        if self.current_arrow_id:
            self.canvas.coords(self.current_arrow_id, self.start_x, self.start_y, event.x, event.y)
    
    def on_button_release(self, event):
        if self.current_arrow_id:
            self.arrow_end = (event.x, event.y)
            self.prompt_for_text(event.x, event.y)

    def prompt_for_text(self, x, y):
        if self.text_entry:
            self.text_entry.destroy()
            
        self.text_entry = tk.Entry(self.canvas)
        self.text_entry.place(x=x, y=y)
        self.text_entry.focus_set()
        self.text_entry.bind("<Return>", self.on_text_entered)
        self.text_entry.bind("<Escape>", self.cancel_text_entry)

    def on_text_entered(self, event):
        text = self.text_entry.get().strip()
        
        self.draw.line((self.start_x, self.start_y, self.arrow_end[0], self.arrow_end[1]), fill="red", width=4)

        if text:
            self.annotations.append(text)
            x, y = self.arrow_end
            try:
                font = ImageFont.truetype("arial.ttf", 20)
            except IOError:
                font = ImageFont.load_default()
            
            text_bbox = self.draw.textbbox((x, y), text, font=font)
            text_width = text_bbox[2] - text_bbox[0]
            text_height = text_bbox[3] - text_bbox[1]
            self.draw.rectangle([x, y, x + text_width + 8, y + text_height + 4], fill="red")
            self.draw.text((x + 4, y + 2), text, fill="white", font=font)

        self.text_entry.destroy()
        self.text_entry = None
        
        self.canvas.delete(self.current_arrow_id)
        self.current_arrow_id = None
        self.update_canvas_image()

    def cancel_text_entry(self, event):
        if self.text_entry:
            self.text_entry.destroy()
            self.text_entry = None
        if self.current_arrow_id:
            self.canvas.delete(self.current_arrow_id)
            self.current_arrow_id = None
            
    def update_canvas_image(self):
        self.photo = ImageTk.PhotoImage(self.annotated_image)
        self.canvas.itemconfig(self.canvas_image, image=self.photo)
    
    def save_annotation(self):
        annotated_filename = f"isaretedilen_{time.time():.2f}.png"
        try:
            self.annotated_image.save(annotated_filename)
            os.remove(self.screenshot_path)
        except Exception as e:
            messagebox.showerror("Hata", f"Anotasyon kaydedilemedi: {e}", **self.messagebox_options)
            self.on_close()
            return
        
        description = "İşaretçi ile tespit."
        if self.annotations:
            description = ", ".join(self.annotations)

        event_info = {
            "type": "screenshot",
            "time": self.current_time_sec,
            "description": description,
            "image": annotated_filename,
            "video": self.parent.current_video
        }
        self.parent.event_list.append(event_info)
        
        # Update the listbox more reliably
        self.parent.events_listbox.delete(0, tk.END)
        for i, event in enumerate(self.parent.event_list):
            desc = event['description']
            time_str = format_time(event['time'])
            self.parent.events_listbox.insert(tk.END, f"Tespit {i + 1} - Zaman: {time_str} - {desc}")
        
        messagebox.showinfo("Bilgi", "İşaretçi tespiti eklendi.", parent=self.parent)
        
        self.on_close(cleanup=False)

    def on_close(self, cleanup=True):
        if cleanup and os.path.exists(self.screenshot_path):
            os.remove(self.screenshot_path)

        if self.was_playing:
            self.parent.play_video()
        
        self.destroy()

if __name__ == "__main__":
    app = VideoAnnotationApp()
    app.mainloop()